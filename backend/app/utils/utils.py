import os
import sys
import openai
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt

from config.config import OUTPUT_DIR


def translate_text(text):
    from openai import AzureOpenAI
    from dotenv import load_dotenv

    load_dotenv()

    client = AzureOpenAI(
        api_version="2024-02-15-preview",
        azure_endpoint="https://swisshacks-aoai-westeurope.openai.azure.com/",
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Translate from German to English."},
            {"role": "user", "content": text},
        ],
    )

    return response.choices[0].message.content


def save_translated_docx(docx_path):
    doc = Document(docx_path)
    translated_doc = Document()
    translated_texts = []
    try:
        for para in doc.paragraphs:
            if para.text.strip():
                translated_text = translate_text(para.text.strip())
                translated_texts.append(translated_text)
                new_para = translated_doc.add_paragraph(translated_text)
                new_para.runs[0].font.size = Pt(11)
    except Exception as e:
        print(f"Error translating text: {e}")
        for translated_text in translated_texts:
            new_para = translated_doc.add_paragraph(translated_text)
            new_para.runs[0].font.size = Pt(11)
    translated_doc.save(f"data/product_portfolio_translated.docx")
